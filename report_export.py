# report_export.py - 专业风资源评估报告导出模块（智能列名识别版）
import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import fit_weibull, weibull_energy_density
from wind_grade import calc_wind_power_density, get_wind_grade


# ══════════════════════════════════════════════
#  智能识别列名
# ══════════════════════════════════════════════
def detect_columns(df):
    """自动识别时间列和风速列"""
    cols = df.columns.tolist()
    
    # 识别时间列
    time_col = None
    for possible in ['datetime', 'time', 'timestamp', 'date', 'Time', 'DateTime']:
        if possible in cols:
            time_col = possible
            break
    
    # 如果列名都找不到，检查索引是否是时间
    if time_col is None:
        if isinstance(df.index, pd.DatetimeIndex):
            df = df.reset_index()
            time_col = df.columns[0]
    
    # 识别风速列
    ws_col = None
    for possible in ['wind_speed', 'ws', 'speed', 'Wind Speed', 'WindSpeed', 'v']:
        if possible in cols:
            ws_col = possible
            break
    
    # 如果还是找不到，尝试第一个数值列
    if ws_col is None:
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        if len(numeric_cols) > 0:
            ws_col = numeric_cols[0]
    
    return time_col, ws_col


# ══════════════════════════════════════════════
#  报告数据汇总计算
# ══════════════════════════════════════════════
def collect_report_data(df, location_name="未命名场址", lat=None, lon=None,
                        height=70, rho=1.225, rated_power=3.0):
    
    # 智能识别列名
    time_col, ws_col = detect_columns(df)
    
    if ws_col is None:
        raise ValueError("❌ 无法识别风速列，请确保数据中包含风速数据")
    
    ws = df[ws_col].dropna().values
    
    # 基础统计
    mean_ws  = round(float(np.mean(ws)), 2)
    std_ws   = round(float(np.std(ws)),  2)
    max_ws   = round(float(np.max(ws)),  2)
    min_ws   = round(float(np.min(ws)),  2)
    
    # 威布尔拟合
    k_fit, c_fit = fit_weibull(ws)
    wpd_actual   = calc_wind_power_density(ws, rho)
    wpd_weibull  = weibull_energy_density(k_fit, c_fit, rho)
    
    # 风资源等级
    grade_info = get_wind_grade(wpd_actual)
    
    # 有效风速
    mask_eff  = (ws >= 3.0) & (ws <= 25.0)
    eff_hours = int(mask_eff.sum())
    eff_ratio = round(eff_hours / len(ws) * 100, 2)
    
    # 月度统计（如果有时间列）
    monthly_avg = None
    data_period = "N/A"
    if time_col is not None:
        try:
            df2 = df.copy()
            df2['_month'] = pd.to_datetime(df2[time_col]).dt.month
            monthly_avg = df2.groupby('_month')[ws_col].mean().round(2)
            
            dt_series = pd.to_datetime(df2[time_col])
            data_period = f"{dt_series.min()} 至 {dt_series.max()}"
        except:
            pass
    
    # 如果月度统计失败，创建默认值
    if monthly_avg is None:
        monthly_avg = pd.Series([mean_ws] * 12, index=range(1, 13))
    
    # 风向统计
    dominant_dir = "N/A"
    wd_col = None
    for possible in ['wind_direction', 'wd', 'direction', 'Wind Direction']:
        if possible in df.columns:
            wd_col = possible
            break
    
    if wd_col is not None:
        wd = df[wd_col].dropna()
        if len(wd) > 0:
            try:
                dir_bins = pd.cut(wd, bins=16, labels=[
                    'N','NNE','NE','ENE','E','ESE','SE','SSE',
                    'S','SSW','SW','WSW','W','WNW','NW','NNW'
                ])
                if len(dir_bins.mode()) > 0:
                    dominant_dir = dir_bins.mode()[0]
            except:
                pass
    
    # 发电量估算
    capacity_factor = min(wpd_actual / 400 * 0.35, 0.45)
    annual_energy   = rated_power * 8760 * capacity_factor
    
    return {
        "项目信息": {
            "场址名称": location_name,
            "纬度":     lat if lat else "N/A",
            "经度":     lon if lon else "N/A",
            "评估高度": f"{height} m",
            "数据时段": data_period,
            "数据量":   f"{len(df):,} 条记录",
        },
        "基础指标": {
            "年平均风速":   f"{mean_ws} m/s",
            "风速标准差":   f"{std_ws} m/s",
            "最大风速":     f"{max_ws} m/s",
            "最小风速":     f"{min_ws} m/s",
            "威布尔 k":     round(k_fit, 3),
            "威布尔 c":     round(c_fit, 2),
        },
        "风能指标": {
            "实测风功率密度":   f"{wpd_actual:.1f} W/m²",
            "威布尔风功率密度": f"{wpd_weibull:.1f} W/m²",
            "国标等级":         f"{grade_info['等级']} - {grade_info['资源描述']}",
            "有效风速小时数":   f"{eff_hours:,} h",
            "有效风速占比":     f"{eff_ratio}%",
        },
        "风向特征": {
            "主导风向": dominant_dir,
        },
        "发电估算": {
            "装机容量":     f"{rated_power} MW",
            "容量系数":     f"{capacity_factor*100:.1f}%",
            "年发电量":     f"{annual_energy:.1f} MWh",
            "等效满发小时": f"{int(annual_energy / rated_power)} h",
        },
        "原始数据": {
            "df":          df,
            "ws":          ws,
            "monthly_avg": monthly_avg,
        },
    }


# ══════════════════════════════════════════════
#  导出 Excel
# ══════════════════════════════════════════════
def export_excel(report_data):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        
        # 工作表1：评估摘要
        records = []
        for section, items in report_data.items():
            if section == "原始数据":
                continue
            for key, val in items.items():
                records.append({"章节": section, "指标": key, "结果": val})
        pd.DataFrame(records).to_excel(writer, sheet_name="评估摘要", index=False)
        
        # 工作表2：月度风速
        monthly = report_data["原始数据"]["monthly_avg"]
        pd.DataFrame({
            "月份": monthly.index,
            "平均风速 (m/s)": monthly.values,
        }).to_excel(writer, sheet_name="月度风速", index=False)
        
        # 工作表3：原始数据样本
        report_data["原始数据"]["df"].head(10000).to_excel(
            writer, sheet_name="原始数据样本", index=False
        )
    
    output.seek(0)
    return output


# ══════════════════════════════════════════════
#  导出 Word
# ══════════════════════════════════════════════
def export_word(report_data):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    
    # 封面
    t = doc.add_heading('风资源评估专业报告', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph(f"项目名称：{report_data['项目信息']['场址名称']}", style='Heading 2')
    doc.add_paragraph(f"评估日期：{datetime.now().strftime('%Y年%m月%d日')}")
    doc.add_paragraph("评估单位：风资源智能评估系统")
    doc.add_page_break()
    
    # 第一章
    doc.add_heading('第一章  项目概况', 1)
    for k, v in report_data["项目信息"].items():
        doc.add_paragraph(f"{k}：{v}", style='List Bullet')
    
    # 第二章
    doc.add_heading('第二章  基础统计指标', 1)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    tbl.rows[0].cells[0].text = '指标'
    tbl.rows[0].cells[1].text = '结果'
    for k, v in report_data["基础指标"].items():
        r = tbl.add_row().cells
        r[0].text = k
        r[1].text = str(v)
    
    # 第三章
    doc.add_heading('第三章  风速特征分析', 1)
    doc.add_paragraph("3.1 月度平均风速")
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.style = 'Light Grid Accent 1'
    tbl2.rows[0].cells[0].text = '月份'
    tbl2.rows[0].cells[1].text = '平均风速 (m/s)'
    for m, v in report_data["原始数据"]["monthly_avg"].items():
        r = tbl2.add_row().cells
        r[0].text = f"{m}月"
        r[1].text = f"{v:.2f}"
    
    # 第四章
    doc.add_heading('第四章  风向特征分析', 1)
    doc.add_paragraph(f"主导风向：{report_data['风向特征']['主导风向']}")
    
    # 第五章
    doc.add_heading('第五章  威布尔拟合结果', 1)
    doc.add_paragraph(f"形状参数 k = {report_data['基础指标']['威布尔 k']}")
    doc.add_paragraph(f"尺度参数 c = {report_data['基础指标']['威布尔 c']}")
    
    # 第六章
    doc.add_heading('第六章  风功率密度等级', 1)
    for k, v in report_data["风能指标"].items():
        doc.add_paragraph(f"{k}：{v}", style='List Bullet')
    
    # 第七章
    doc.add_heading('第七章  发电量估算', 1)
    for k, v in report_data["发电估算"].items():
        doc.add_paragraph(f"{k}：{v}", style='List Bullet')
    
    # 第八章
    doc.add_heading('第八章  综合评价结论', 1)
    wpd = float(report_data["风能指标"]["实测风功率密度"].split()[0])
    if wpd >= 200:
        conclusion = (
            f"该场址年平均风速为 {report_data['基础指标']['年平均风速']}，"
            f"实测风功率密度为 {report_data['风能指标']['实测风功率密度']}，"
            f"达到国标 {report_data['风能指标']['国标等级']} 标准，"
            f"具备商业化风电开发条件，建议进一步开展微观选址和详细设计工作。"
        )
    else:
        conclusion = (
            f"该场址年平均风速为 {report_data['基础指标']['年平均风速']}，"
            f"实测风功率密度为 {report_data['风能指标']['实测风功率密度']}，"
            f"属于国标 {report_data['风能指标']['国标等级']} 标准，"
            f"当前风能资源条件尚不满足大规模商业化开发要求，"
            f"建议延长观测周期或考虑更换场址。"
        )
    doc.add_paragraph(conclusion)
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# ══════════════════════════════════════════════
#  Streamlit 页面
# ══════════════════════════════════════════════
def show_report_export_page(df=None, rho=1.225, height=70, key_prefix="report_"):
    st.markdown("## 📄 专业评估报告导出")
    st.caption("生成符合行业规范的完整风资源评估报告，可直接用于工程决策")
    
    if df is None or len(df) == 0:
        st.warning("⚠️ 请先在左侧生成或上传数据后再使用本模块。")
        return
    
    # 显示检测到的列名（调试信息）
    time_col, ws_col = detect_columns(df)
    with st.expander("🔍 数据列检测结果", expanded=False):
        st.write(f"✅ 风速列：`{ws_col}`")
        st.write(f"✅ 时间列：`{time_col if time_col else '未检测到（将使用默认值）'}`")
        st.write(f"📋 所有列名：{df.columns.tolist()}")
    
    st.divider()
    
    # ── 参数设置 ──────────────────────────────
    st.subheader("⚙️ 报告参数设置")
    col1, col2, col3 = st.columns(3)
    with col1:
        location_name = st.text_input("场址名称", value="示例风电场",
                                      key=f"{key_prefix}location")
    with col2:
        lat = st.number_input("纬度", value=30.0, format="%.4f",
                              key=f"{key_prefix}lat")
    with col3:
        lon = st.number_input("经度", value=120.0, format="%.4f",
                              key=f"{key_prefix}lon")
    
    col4, col5, col6 = st.columns(3)
    with col4:
        height_input = st.number_input("评估高度 (m)", value=int(height),
                                       min_value=10, max_value=200,
                                       key=f"{key_prefix}height")
    with col5:
        rho_input = st.number_input("空气密度 (kg/m³)", value=float(rho),
                                    format="%.3f", key=f"{key_prefix}rho")
    with col6:
        rated_power = st.number_input("装机容量 (MW)", value=3.0,
                                      min_value=0.1, max_value=100.0,
                                      key=f"{key_prefix}power")
    
    st.divider()
    
    # ── 汇总数据 ──────────────────────────────
    try:
        report_data = collect_report_data(
            df, location_name, lat, lon,
            height_input, rho_input, rated_power
        )
    except Exception as e:
        st.error(f"❌ 数据处理错误：{str(e)}")
        st.stop()
    
    # ── 报告预览 ──────────────────────────────
    st.subheader("📋 报告内容预览")
    tabs = st.tabs([
        "📍 项目概况", "📊 基础指标", "🌬️ 风速特征",
        "🧭 风向特征", "📈 威布尔拟合", "⚡ 风能等级",
        "💡 发电估算", "✅ 评价结论"
    ])
    
    with tabs[0]:
        for k, v in report_data["项目信息"].items():
            st.write(f"**{k}**：{v}")
    
    with tabs[1]:
        for k, v in report_data["基础指标"].items():
            st.write(f"**{k}**：{v}")
    
    with tabs[2]:
        monthly = report_data["原始数据"]["monthly_avg"]
        st.bar_chart(monthly)
        st.caption("各月平均风速 (m/s)")
    
    with tabs[3]:
        st.info(f"主导风向：{report_data['风向特征']['主导风向']}")
    
    with tabs[4]:
        c1, c2 = st.columns(2)
        c1.metric("威布尔形状参数 k", report_data["基础指标"]["威布尔 k"])
        c2.metric("威布尔尺度参数 c", report_data["基础指标"]["威布尔 c"])
    
    with tabs[5]:
        for k, v in report_data["风能指标"].items():
            st.write(f"**{k}**：{v}")
    
    with tabs[6]:
        cols = st.columns(2)
        for i, (k, v) in enumerate(report_data["发电估算"].items()):
            cols[i % 2].metric(k, v)
    
    with tabs[7]:
        wpd = float(report_data["风能指标"]["实测风功率密度"].split()[0])
        if wpd >= 200:
            st.success(
                f"✅ 该场址风功率密度达到 **{wpd:.1f} W/m²**，"
                f"**具备商业化开发条件**，建议进一步开展详细设计。"
            )
        else:
            st.warning(
                f"⚠️ 该场址风功率密度为 **{wpd:.1f} W/m²**，"
                f"**暂不满足商业化开发要求**（建议 ≥ 200 W/m²）。"
            )
    
    st.divider()
    
    # ── 导出按钮 ──────────────────────────────
    st.subheader("⬇️ 导出报告")
    col_a, col_b, col_c = st.columns(3)
    
    with col_a:
        excel_bytes = export_excel(report_data)
        st.download_button(
            label="📥 下载 Excel 汇总表",
            data=excel_bytes,
            file_name=f"{location_name}_评估报告_{datetime.now().strftime('%Y%m%d')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key=f"{key_prefix}dl_excel",
            use_container_width=True,
        )
    
    with col_b:
        word_bytes = export_word(report_data)
        st.download_button(
            label="📥 下载 Word 专业报告",
            data=word_bytes,
            file_name=f"{location_name}_评估报告_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"{key_prefix}dl_word",
            use_container_width=True,
        )
    
    with col_c:
        st.info("💡 PDF 格式：\n打开下载的 Word 文件\n→ 文件 → 另存为 → PDF", icon="ℹ️")
    
    st.divider()
    
    # ── 使用说明 ──────────────────────────────
    with st.expander("📖 报告使用说明 / 答辩演示建议"):
        st.markdown("""
        ### 报告章节说明
        | 章节 | 内容 |
        |------|------|
        | 第一章 | 场址基本信息、数据来源与时段 |
        | 第二章 | 风速统计量、威布尔参数 |
        | 第三章 | 月度风速变化规律 |
        | 第四章 | 主导风向特征 |
        | 第五章 | 威布尔分布拟合结果 |
        | 第六章 | 国标风功率密度等级评定 |
        | 第七章 | 容量系数与年发电量估算 |
        | 第八章 | 综合评价与开发建议 |
        
        ### 导出格式说明
        - 📊 **Excel**：评估摘要 + 月度数据 + 原始数据样本（三张工作表）
        - 📄 **Word**：完整专业报告，含所有章节和数据表格
        - 📋 **PDF**：将 Word 另存为 PDF 即可
        
        ### 🎤 答辩演示建议
        1. 在左侧侧边栏生成数据
        2. 点击「📄 报告导出」tab
        3. 填写场址名称和参数
        4. 逐一点击预览各章节内容
        5. 点击「下载 Word 专业报告」
        6. 打开 Word 文件展示专业排版效果
        """)

